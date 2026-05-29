# File: biomedical_experiment_analysis_system.py

"""
Biomedical Experiment Analysis System
-------------------------------------
A Python framework to load experimental data, run basic statistics,
generate publication-quality figures, and export PDF/Word lab reports.
"""

import os
import pandas as pd
import numpy as np
from scipy import stats
import matplotlib.pyplot as plt
import seaborn as sns
from matplotlib.backends.backend_pdf import PdfPages
from docx import Document
from docx.shared import Inches

sns.set(style="whitegrid")

# -----------------------------
# Data Loader
# -----------------------------
class ExperimentDataLoader:
    """
    Loads experimental CSV/XLSX data with flexible group columns.
    Assumes first column is 'Sample' or identifier; all others are groups.
    Converts wide format to long format: columns = ['Sample', 'Group', 'Measurement']
    """
    def __init__(self, filepath: str):
        self.filepath = filepath
        self.data = None

    def load(self) -> pd.DataFrame:
        # Load file
        if self.filepath.endswith(".csv"):
            df = pd.read_csv(self.filepath)
        elif self.filepath.endswith((".xls", ".xlsx")):
            df = pd.read_excel(self.filepath)
        else:
            raise ValueError("Unsupported file format. Use CSV or Excel.")

        if df.shape[1] < 2:
            raise ValueError("Data must have at least one sample column and one group column.")

        # Automatically reshape to long format
        sample_col = df.columns[0]
        group_cols = df.columns[1:]
        df_long = df.melt(id_vars=sample_col, value_vars=group_cols,
                          var_name="Group", value_name="Measurement")
        df_long.rename(columns={sample_col: "Sample"}, inplace=True)

        self.data = df_long
        return self.data

# -----------------------------
# Statistical Analysis
# -----------------------------
class StatisticalAnalyzer:
    """
    Computes descriptive statistics and basic inferential tests.
    """
    def __init__(self, data: pd.DataFrame):
        self.data = data
        self.results = {}

    def descriptive_stats(self) -> pd.DataFrame:
        desc = self.data.groupby("Group")["Measurement"].agg(
            ["mean", "std", "count"]
        )
        desc["sem"] = desc["std"] / np.sqrt(desc["count"])
        self.results["descriptive"] = desc
        return desc

    def t_test(self, group1: str, group2: str) -> dict:
        data1 = self.data[self.data["Group"] == group1]["Measurement"]
        data2 = self.data[self.data["Group"] == group2]["Measurement"]
        t_stat, p_val = stats.ttest_ind(data1, data2)
        self.results[f"t_{group1}_vs_{group2}"] = {"t": t_stat, "p": p_val}
        return {"t": t_stat, "p": p_val}

    def run_all_tests(self) -> dict:
        self.descriptive_stats()
        groups = self.data["Group"].unique()
        if len(groups) >= 2:
            for i in range(len(groups)):
                for j in range(i + 1, len(groups)):
                    self.t_test(groups[i], groups[j])
        return self.results

# -----------------------------
# Figure Generator
# -----------------------------
class FigureGenerator:
    """
    Generates high-resolution figures for each group.
    """
    def __init__(self, data: pd.DataFrame, output_dir: str = "figures"):
        self.data = data
        self.output_dir = output_dir
        os.makedirs(self.output_dir, exist_ok=True)
        self.figures = []

    def boxplot(self) -> str:
        plt.figure(figsize=(6, 4))
        sns.boxplot(x="Group", y="Measurement", data=self.data, palette="Set2")
        plt.title("Experimental Measurements by Group")
        plt.tight_layout()
        path = os.path.join(self.output_dir, "boxplot.png")
        plt.savefig(path, dpi=300)
        plt.close()
        self.figures.append(path)
        return path

    def barplot(self) -> str:
        means = self.data.groupby("Group")["Measurement"].mean().reset_index()
        plt.figure(figsize=(6, 4))
        sns.barplot(x="Group", y="Measurement", data=means, palette="Set2")
        plt.title("Mean Measurement by Group")
        plt.tight_layout()
        path = os.path.join(self.output_dir, "barplot.png")
        plt.savefig(path, dpi=300)
        plt.close()
        self.figures.append(path)
        return path

# -----------------------------
# Report Generator
# -----------------------------
class ReportGenerator:
    """
    Generates PDF and Word reports with figures and statistical summaries.
    """
    def __init__(self, stats: dict, figures: list, title: str = "Lab Report"):
        self.stats = stats
        self.figures = figures
        self.title = title

    def export_pdf(self, output_file: str):
        with PdfPages(output_file) as pdf:
            for fig_path in self.figures:
                fig = plt.imread(fig_path)
                plt.figure(figsize=(6, 4))
                plt.imshow(fig)
                plt.axis("off")
                pdf.savefig()
                plt.close()
        print(f"PDF report saved: {output_file}")

    def export_word(self, output_file: str):
        doc = Document()
        doc.add_heading(self.title, level=0)

        # Add descriptive statistics
        if "descriptive" in self.stats:
            doc.add_heading("Descriptive Statistics", level=1)
            desc = self.stats["descriptive"]
            table = doc.add_table(rows=1, cols=len(desc.columns) + 1)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Group"
            for i, col in enumerate(desc.columns):
                hdr_cells[i + 1].text = col
            for group, row in desc.iterrows():
                cells = table.add_row().cells
                cells[0].text = str(group)
                for i, col in enumerate(desc.columns):
                    cells[i + 1].text = f"{row[col]:.3f}"

        # Add figures
        doc.add_heading("Figures", level=1)
        for fig_path in self.figures:
            doc.add_picture(fig_path, width=Inches(5))
        doc.save(output_file)
        print(f"Word report saved: {output_file}")

# -----------------------------
# Example Workflow
# -----------------------------
if __name__ == "__main__":
    # Hardcoded file path
    filepath = r"C:\Users\USER\Documents\Glucose data(sample).csv"

    # Load data
    loader = ExperimentDataLoader(filepath)
    data = loader.load()
    print(f"Data loaded successfully. Columns: {data.columns.tolist()}")
    print(data.head())

    # Run analysis
    analyzer = StatisticalAnalyzer(data)
    stats_results = analyzer.run_all_tests()
    print("Statistical analysis complete.")

    # Generate figures
    fig_gen = FigureGenerator(data)
    fig_gen.boxplot()
    fig_gen.barplot()
    print(f"Figures saved in: {fig_gen.output_dir}")

    # Export reports
    report = ReportGenerator(stats_results, fig_gen.figures)
    report.export_pdf("lab_report.pdf")
    report.export_word("lab_report.docx")