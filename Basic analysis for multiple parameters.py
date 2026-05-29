# File: biomedical_experiment_multi_parameter_bar_chart.py

import os
import pandas as pd
import numpy as np
from scipy import stats
import matplotlib.pyplot as plt
import seaborn as sns
from datetime import datetime
from docx import Document
from docx.shared import Inches

sns.set(style="whitegrid")

# -----------------------------
# Data Loader
# -----------------------------
class ExperimentDataLoader:
    def __init__(self, filepath: str, sample_col: str = 'Animal ID', group_col: str = 'Group'):
        self.filepath = filepath
        self.data = None
        self.sample_col = sample_col
        self.group_col = group_col

    def load(self) -> pd.DataFrame:
        df = pd.read_csv(self.filepath)
        self.data = df.rename(columns={self.sample_col: 'Sample', self.group_col: 'Group'})
        self.sample_col = 'Sample'
        self.group_col = 'Group'
        return self.data

    def get_numeric_parameters(self):
        numeric_cols = []
        for col in self.data.columns[2:]:  # exclude Sample and Group columns
            self.data[col] = pd.to_numeric(self.data[col], errors='coerce')
            if self.data[col].notna().any():
                numeric_cols.append(col)
        return numeric_cols

# -----------------------------
# Statistical Analysis
# -----------------------------
class MultiParameterAnalyzer:
    def __init__(self, data: pd.DataFrame, group_col: str):
        self.data = data
        self.group_col = group_col
        self.results_text = ""

    def analyze_parameter(self, parameter: str):
        df = self.data[[self.group_col, parameter]].dropna()

        stats_df = df.groupby(self.group_col)[parameter].agg(['mean','std','count']).reset_index()
        stats_df.rename(columns={'mean':'Mean','std':'SD','count':'n'}, inplace=True)
        stats_df['Parameter'] = parameter

        groups = df[self.group_col].unique()
        comparisons = []
        for i in range(len(groups)):
            for j in range(i+1, len(groups)):
                g1 = df[df[self.group_col]==groups[i]][parameter]
                g2 = df[df[self.group_col]==groups[j]][parameter]
                t_stat, p_val = stats.ttest_ind(g1, g2, nan_policy='omit')
                comparisons.append({'Parameter': parameter, 'Group1': groups[i], 'Group2': groups[j], 't': t_stat, 'p': p_val})
                self.results_text += f"{parameter} {groups[i]} vs {groups[j]}: t={t_stat:.3f}, p={p_val:.4f}\n"

        comparisons_df = pd.DataFrame(comparisons)
        return stats_df, comparisons_df

# -----------------------------
# Figure Generator (Bar Charts)
# -----------------------------
class FigureGenerator:
    def __init__(self, data: pd.DataFrame, group_col: str, output_dir: str):
        self.data = data
        self.group_col = group_col
        self.output_dir = output_dir
        os.makedirs(self.output_dir, exist_ok=True)

    def plot_parameter(self, parameter: str) -> str:
        plot_data = self.data[[self.group_col, parameter]].dropna()
        means = plot_data.groupby(self.group_col)[parameter].mean().reset_index()

        plt.figure(figsize=(6,4))
        sns.barplot(x=self.group_col, y=parameter, data=means, palette='Set2')
        plt.title(f'{parameter} by {self.group_col}')
        plt.tight_layout()
        path = os.path.join(self.output_dir, f'{parameter}_barchart.png')
        plt.savefig(path, dpi=300)
        plt.close()
        return path

# -----------------------------
# Word + CSV Report Generator
# -----------------------------
class ReportGenerator:
    def __init__(self, stats_df: pd.DataFrame, comparisons_df: pd.DataFrame, results_text: str, output_dir: str):
        self.stats_df = stats_df
        self.comparisons_df = comparisons_df
        self.results_text = results_text
        self.output_dir = output_dir
        os.makedirs(self.output_dir, exist_ok=True)

    def export_word(self, filename="lab_report.docx"):
        path = os.path.join(self.output_dir, filename)
        doc = Document()
        doc.add_heading("Lab Report", level=0)

        doc.add_heading("Summary Statistics", level=1)
        params = self.stats_df['Parameter'].unique()
        for param in params:
            df = self.stats_df[self.stats_df['Parameter']==param]
            doc.add_heading(param, level=2)
            table = doc.add_table(rows=1, cols=len(df.columns))
            hdr_cells = table.rows[0].cells
            for i, col in enumerate(df.columns):
                hdr_cells[i].text = col
            for _, row in df.iterrows():
                cells = table.add_row().cells
                for i, col in enumerate(df.columns):
                    cells[i].text = str(row[col])

        doc.add_heading("Comparisons (t-tests)", level=1)
        doc.add_paragraph(self.results_text)
        doc.save(path)
        return path

    def export_csv(self):
        stats_path = os.path.join(self.output_dir, "summary_statistics.csv")
        comps_path = os.path.join(self.output_dir, "summary_comparisons.csv")
        self.stats_df.to_csv(stats_path, index=False)
        self.comparisons_df.to_csv(comps_path, index=False)
        return stats_path, comps_path

# -----------------------------
# Main Workflow
# -----------------------------
if __name__ == "__main__":
    filepath = r"C:\Users\USER\Documents\Python analysis\Multiple glucose parameters.csv"
    loader = ExperimentDataLoader(filepath)
    data = loader.load()
    data.head()

    numeric_params = loader.get_numeric_parameters()
    print(f"Numeric parameters detected: {numeric_params}")

    base_folder = os.path.dirname(filepath)
    experiment_name = os.path.splitext(os.path.basename(filepath))[0]
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    results_folder = os.path.join(base_folder, f'results_{experiment_name}_{timestamp}')
    os.makedirs(results_folder, exist_ok=True)

    all_stats = []
    all_comparisons = []
    fig_gen = FigureGenerator(data, group_col='Group', output_dir=results_folder)

    for param in numeric_params:
        analyzer = MultiParameterAnalyzer(data, group_col='Group')
        stats_df, comparisons_df = analyzer.analyze_parameter(param)
        all_stats.append(stats_df)
        all_comparisons.append(comparisons_df)
        fig_gen.plot_parameter(param)

    summary_stats_df = pd.concat(all_stats, ignore_index=True)
    summary_comparisons_df = pd.concat(all_comparisons, ignore_index=True)

    report = ReportGenerator(summary_stats_df, summary_comparisons_df, analyzer.results_text, results_folder)
    report.export_csv()
    report.export_word()

    print(f"All results saved in folder: {results_folder}")