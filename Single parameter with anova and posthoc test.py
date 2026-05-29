# File: biomedical_experiment_analysis_system.py

"""
Biomedical Experiment Analysis System
-------------------------------------
A Python framework to load experimental data, run basic statistics,
generate publication-quality figures, and export PDF/Word lab reports.
Includes ANOVA + post-hoc and auto-generated results text.
"""
import os
import pandas as pd
import numpy as np
from scipy import stats
import statsmodels.api as sm
from statsmodels.formula.api import ols
import matplotlib.pyplot as plt
import seaborn as sns
from matplotlib.backends.backend_pdf import PdfPages
from docx import Document
from docx.shared import Inches

sns.set(style="whitegrid")

# -----------------------------
# Data Loader
# -----------------------------
class ExperimentDataLoader:
    """
    Loads experimental CSV/XLSX data with flexible group columns.
    Converts wide format to long format: columns = ['Sample', 'Group', 'Measurement']
    """
    def __init__(self, filepath: str):
        self.filepath = filepath
        self.data = None

    def load(self) -> pd.DataFrame:
        if self.filepath.endswith(".csv"):
            df = pd.read_csv(self.filepath)
        elif self.filepath.endswith((".xls", ".xlsx")):
            df = pd.read_excel(self.filepath)
        else:
            raise ValueError("Unsupported file format. Use CSV or Excel.")

        if df.shape[1] < 2:
            raise ValueError("Data must have at least one sample column and one group column.")

        # Reshape wide -> long format
        sample_col = df.columns[0]
        group_cols = df.columns[1:]
        df_long = df.melt(id_vars=sample_col, value_vars=group_cols,
                          var_name="Group", value_name="Measurement")
        df_long.rename(columns={sample_col: "Sample"}, inplace=True)

        self.data = df_long
        return self.data

# -----------------------------
# Statistical Analysis
# -----------------------------
class StatisticalAnalyzer:
    """
    Computes descriptive statistics, t-tests, ANOVA + post-hoc,
    and generates results text with significance.
    """
    def __init__(self, data: pd.DataFrame):
        self.data = data
        self.results = {}
        self.results_text = ""

    def descriptive_stats(self) -> pd.DataFrame:
        desc = self.data.groupby("Group")["Measurement"].agg(["mean", "std", "count"])
        desc["sem"] = desc["std"] / np.sqrt(desc["count"])
        self.results["descriptive"] = desc
        return desc

    def t_test(self, group1: str, group2: str) -> dict:
        data1 = self.data[self.data["Group"] == group1]["Measurement"]
        data2 = self.data[self.data["Group"] == group2]["Measurement"]
        t_stat, p_val = stats.ttest_ind(data1, data2)
        self.results[f"t_{group1}_vs_{group2}"] = {"t": t_stat, "p": p_val}
        self.results_text += f"t-test {group1} vs {group2}: t={t_stat:.3f}, p={p_val:.4f}\n"
        return {"t": t_stat, "p": p_val}

    def anova_posthoc(self):
        # One-way ANOVA
        model = ols('Measurement ~ C(Group)', data=self.data).fit()
        anova_table = sm.stats.anova_lm(model, typ=2)
        self.results["ANOVA"] = anova_table
        self.results_text += f"\nOne-way ANOVA:\n{anova_table}\n"

        # Post-hoc Tukey HSD
        from statsmodels.stats.multicomp import pairwise_tukeyhsd
        tukey = pairwise_tukeyhsd(endog=self.data['Measurement'], groups=self.data['Group'], alpha=0.05)
        self.results["TukeyHSD"] = tukey.summary()
        self.results_text += f"\nTukey HSD post-hoc results:\n{tukey.summary()}\n"

    def run_all_tests(self) -> dict:
        self.descriptive_stats()
        groups = self.data["Group"].unique()

        # Pairwise t-tests if 2 groups
        if len(groups) == 2:
            self.t_test(groups[0], groups[1])

        # ANOVA + post-hoc if more than 2 groups
        if len(groups) > 2:
            self.anova_posthoc()

        return self.results

# -----------------------------
# Figure Generator
# -----------------------------
class FigureGenerator:
    """
    Generates high-resolution figures for each group.
    """
    def __init__(self, data: pd.DataFrame, output_dir: str):
        self.data = data
        self.output_dir = output_dir
        os.makedirs(self.output_dir, exist_ok=True)
        self.figures = []

    def boxplot(self) -> str:
        plt.figure(figsize=(6, 4))
        sns.boxplot(x="Group", y="Measurement", data=self.data, palette="Set2")
        plt.title("Experimental Measurements by Group")
        plt.tight_layout()
        path = os.path.join(self.output_dir, "boxplot.png")
        plt.savefig(path, dpi=300)
        plt.close()
        self.figures.append(path)
        return path

    def barplot(self) -> str:
        means = self.data.groupby("Group")["Measurement"].mean().reset_index()
        plt.figure(figsize=(6, 4))
        sns.barplot(x="Group", y="Measurement", data=means, palette="Set2")
        plt.title("Mean Measurement by Group")
        plt.tight_layout()
        path = os.path.join(self.output_dir, "barplot.png")
        plt.savefig(path, dpi=300)
        plt.close()
        self.figures.append(path)
        return path

# -----------------------------
# Report Generator
# -----------------------------
class ReportGenerator:
    """
    Generates PDF and Word reports with figures, statistical summaries, and results text.
    """
    def __init__(self, stats: dict, figures: list, results_text: str, output_dir: str, title: str = "Lab Report"):
        self.stats = stats
        self.figures = figures
        self.results_text = results_text
        self.title = title
        self.output_dir = output_dir
        os.makedirs(self.output_dir, exist_ok=True)

    def export_pdf(self, output_file: str):
        path = os.path.join(self.output_dir, output_file)
        with PdfPages(path) as pdf:
            for fig_path in self.figures:
                fig = plt.imread(fig_path)
                plt.figure(figsize=(6, 4))
                plt.imshow(fig)
                plt.axis("off")
                pdf.savefig()
                plt.close()

            # Add a page with results text
            plt.figure(figsize=(6, 4))
            plt.text(0.01, 0.01, self.results_text, fontsize=8, va='bottom')
            plt.axis('off')
            pdf.savefig()
            plt.close()

        print(f"PDF report saved: {path}")

    def export_word(self, output_file: str):
        path = os.path.join(self.output_dir, output_file)
        doc = Document()
        doc.add_heading(self.title, level=0)

        # Add descriptive statistics
        if "descriptive" in self.stats:
            doc.add_heading("Descriptive Statistics", level=1)
            desc = self.stats["descriptive"]
            table = doc.add_table(rows=1, cols=len(desc.columns) + 1)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Group"
            for i, col in enumerate(desc.columns):
                hdr_cells[i + 1].text = col
            for group, row in desc.iterrows():
                cells = table.add_row().cells
                cells[0].text = str(group)
                for i, col in enumerate(desc.columns):
                    cells[i + 1].text = f"{row[col]:.3f}"

        # Add figures
        doc.add_heading("Figures", level=1)
        for fig_path in self.figures:
            doc.add_picture(fig_path, width=Inches(5))

        # Add results text
        doc.add_heading("Results Summary", level=1)
        doc.add_paragraph(self.results_text)

        doc.save(path)
        print(f"Word report saved: {path}")

# -----------------------------
# Example Workflow
# -----------------------------
if __name__ == "__main__":
    # Hardcoded file path
    filepath = r"C:\Users\USER\Documents\Glucose data(sample).csv"

    # Output folder same as input folder
    output_dir = os.path.dirname(filepath)

    # Load data
    loader = ExperimentDataLoader(filepath)
    data = loader.load()
    print(f"Data loaded successfully. Columns: {data.columns.tolist()}")
    print(data.head())

    # Run analysis
    analyzer = StatisticalAnalyzer(data)
    stats_results = analyzer.run_all_tests()
    print("Statistical analysis complete.")

    # Generate figures
    fig_gen = FigureGenerator(data, output_dir)
    fig_gen.boxplot()
    fig_gen.barplot()
    print(f"Figures saved in: {fig_gen.output_dir}")

    # Export reports
    report = ReportGenerator(stats_results, fig_gen.figures, analyzer.results_text, output_dir)
    report.export_pdf("lab_report.pdf")
    report.export_word("lab_report.docx")
